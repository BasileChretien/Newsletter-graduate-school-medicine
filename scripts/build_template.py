"""Build the modernized MERIDIAN newsletter DOCX template.

Reads the original NagoyaU template, preserves all section names and content
verbatim (text only), and restyles it per the Meridian visual spec:

- Title: MERIDIAN (replaces MEDICAL FRONTIER)
- Wine red + warm gold palette
- Cambria headings, Calibri body
- Masthead band, section bars, gold dividers
- Zebra-styled data tables, refined header/footer
"""

from __future__ import annotations

import logging
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from scripts.config import (
    DEAN_NAME,
    DEAN_NAME_PLAIN,
    DEAN_PATH,
    DEAN_PHOTO_PLACEHOLDER,
    DEAN_TITLE,
    LOGO_PATH,
    MERIDIAN_TEMPLATE,
    ORIGINAL_TEMPLATE,
    PALETTE,
    SUBHEAD_TEXTS,
    SUBTITLE,
    TAGLINE,
    TITLE,
)
from scripts.oxml_helpers import (
    add_page_field,
    remove_table_borders,
    set_cell_borders,
    set_cell_margins,
    set_cell_shading,
    set_paragraph_border,
    set_row_height,
    set_run_letter_spacing,
    set_run_small_caps,
    set_table_fixed_layout,
)


# ---------- color helpers ----------
def rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


PRIMARY = rgb(PALETTE["primary"])
ACCENT = rgb(PALETTE["accent"])
ACCENT_AA = rgb(PALETTE["accent_aa"])
TEXT = rgb(PALETTE["text"])
MUTED = rgb(PALETTE["muted"])


# Index of each table inside the DOCX (in document order). Named so the
# orchestration in `build()` is grep-able instead of mystery-meat
# `doc.tables[3]` calls.
TABLE_MASTHEAD = 0
TABLE_DEAN = 1
TABLE_HIGHLIGHTS_TOP = 2
TABLE_HIGHLIGHTS_BOTTOM = 3
TABLE_VISITORS = 4
TABLE_EVENTS = 5
TABLE_CONTACT = 6


def _normalize_body_run(run) -> None:
    """Apply default body styling (Calibri / charcoal / 10.5pt) to a run.

    Used inside table cells to avoid the original template's Arial /
    blue / unsized defaults bleeding through. Idempotent -- non-default
    fields the editor sets are preserved.
    """
    if run.font.name in (None, "", "Arial"):
        run.font.name = "Calibri"
    if run.font.size is None:
        run.font.size = Pt(10.5)
    if run.font.color.rgb is None:
        run.font.color.rgb = TEXT


# ---------- run/paragraph styling helpers ----------
def style_run(run, *, font=None, size_pt=None, bold=None, italic=None,
              color: RGBColor | None = None, all_caps=False, small_caps=False,
              tracking=None) -> None:
    if font:
        run.font.name = font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if all_caps:
        run.font.all_caps = True
    if small_caps:
        set_run_small_caps(run)
    if tracking is not None:
        set_run_letter_spacing(run, tracking)


def style_paragraph(p, *, alignment=None, space_before=None, space_after=None,
                    line_spacing=None, left_indent=None) -> None:
    pf = p.paragraph_format
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)


SECTION_HEAD_RE = re.compile(r"^\s*(\d+)\.\s+(.+?)\s*$")


def is_section_heading(text: str) -> bool:
    """Match patterns like '1.  Message from the Dean / Director'."""
    return bool(SECTION_HEAD_RE.match(text))


# ---------- masthead ----------
def restyle_masthead(table) -> None:
    """Table 0: 1 row x 2 cols. Replace contents with MERIDIAN masthead."""
    remove_table_borders(table)

    # We need a 3-row layout. Strip the existing single-row table content
    # and rebuild it as masthead. Easiest: keep existing table but rebuild
    # the right cell, then insert two extra rows manually.
    # Simpler approach: replace text in the existing 1-row table to be the
    # combined masthead and skip multi-row redesign (limit complexity).

    # We'll use the right-hand cell as the masthead cell.
    cell = table.rows[0].cells[1]
    # Clear existing paragraphs in the cell (keep one).
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    # Cell shading: cream background.
    set_cell_shading(cell, PALETTE["cream"])
    set_cell_margins(cell, top=240, bottom=160, left=200, right=160)
    set_cell_borders(
        cell,
        top={"sz": 48, "color": PALETTE["primary"], "val": "single"},
        bottom={"sz": 8, "color": PALETTE["accent"], "val": "single"},
    )

    # H1 title
    p_title = cell.add_paragraph()
    style_paragraph(p_title, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=0, space_after=2, line_spacing=1.0)
    r = p_title.add_run(TITLE)
    style_run(r, font="Cambria", size_pt=36, bold=True,
              color=PRIMARY, all_caps=True, tracking=20)

    # Tagline (now promoted to 14pt, with a gold underline rule)
    p_tag = cell.add_paragraph()
    style_paragraph(p_tag, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=4, space_after=6, line_spacing=1.2)
    r = p_tag.add_run(TAGLINE)
    style_run(r, font="Cambria", size_pt=14, italic=True, color=PRIMARY)
    set_paragraph_border(p_tag, position="bottom", sz=4,
                         color=PALETTE["accent"], val="single", space=4)

    # Subtitle (now smaller, secondary)
    p_sub = cell.add_paragraph()
    style_paragraph(p_sub, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=4, space_after=8, line_spacing=1.15)
    r = p_sub.add_run(SUBTITLE)
    style_run(r, font="Calibri", size_pt=10, italic=True, color=MUTED)

    # Issue line
    p_iss = cell.add_paragraph()
    style_paragraph(p_iss, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=4, space_after=0)
    parts = [("VOL. XX", PRIMARY), ("  |  ", ACCENT_AA),
             ("ISSUE NO. XX", PRIMARY), ("  |  ", ACCENT_AA),
             ("MONTH YEAR", PRIMARY)]
    for text, color in parts:
        r = p_iss.add_run(text)
        style_run(r, font="Calibri", size_pt=9, bold=True,
                  color=color, small_caps=True, tracking=40)

    # Left cell: NU graduate-school-of-medicine logo on cream backdrop.
    left = table.rows[0].cells[0]
    for p in list(left.paragraphs):
        p._element.getparent().remove(p._element)
    set_cell_shading(left, PALETTE["cream"])
    set_cell_margins(left, top=200, bottom=160, left=160, right=80)
    set_cell_borders(
        left,
        top={"sz": 48, "color": PALETTE["primary"], "val": "single"},
        bottom={"sz": 8, "color": PALETTE["accent"], "val": "single"},
    )
    p_logo = left.add_paragraph()
    style_paragraph(p_logo, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=0, space_after=0)
    if LOGO_PATH.exists():
        run_logo = p_logo.add_run()
        # Reduced from 0.95" to 0.7" to dial back the blue logo's competition
        # with the wine-red MERIDIAN wordmark.
        run_logo.add_picture(str(LOGO_PATH), width=Inches(0.70))
    # Column widths: ~0.9" logo column, ~6.0" masthead column.
    left.width = Inches(0.90)
    cell.width = Inches(6.00)


# ---------- section heading restyle ----------
def restyle_section_heading(p) -> None:
    """Reformat '1.  Message from the Dean / Director' style headings."""
    text = p.text
    m = SECTION_HEAD_RE.match(text)
    if not m:
        return
    num, label = m.group(1), m.group(2)
    # Clear existing runs.
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=18, space_after=6, line_spacing=1.1)
    # Add red leading "01 — " then red label, with the dash in AA gold.
    pad = num.zfill(2)
    r1 = p.add_run(f"{pad} ")
    style_run(r1, font="Cambria", size_pt=16, bold=True,
              color=PRIMARY, all_caps=True, tracking=15)
    r_dash = p.add_run(" —  ")
    style_run(r_dash, font="Cambria", size_pt=16, bold=True,
              color=ACCENT_AA, all_caps=True, tracking=15)
    r2 = p.add_run(label.upper())
    style_run(r2, font="Cambria", size_pt=16, bold=True,
              color=PRIMARY, all_caps=True, tracking=20)
    # Vertical accent bar to the LEFT of the heading text (mirrors the HTML
    # `border-left: 4px solid #8B1A1F`). Word renders paragraph left borders
    # only as tall as the paragraph — so the bar is short, not a full-width
    # underline.
    set_paragraph_border(p, position="left", sz=24, color=PALETTE["primary"],
                         val="single", space=8)
    p.paragraph_format.left_indent = Inches(0.12)


# ---------- subhead (Notable Publications, etc.) ----------
# SUBHEAD_TEXTS imported from scripts.config (single source of truth)


def restyle_subhead(p) -> None:
    """Format paragraph as an H3 sub-heading (Cambria 12pt bold charcoal)."""
    text = p.text
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=10, space_after=4)
    r = p.add_run(text)
    style_run(r, font="Cambria", size_pt=12, bold=True, color=TEXT)


# ---------- body paragraph ----------
def restyle_body(p) -> None:
    """Apply default body styling: Calibri 10.5pt charcoal."""
    if not p.runs:
        return
    style_paragraph(p, space_after=6, line_spacing=1.30)
    for r in p.runs:
        _normalize_body_run(r)


# ---------- bullet list ----------
def restyle_bullet(p) -> None:
    style_paragraph(p, space_after=4, line_spacing=1.25, left_indent=0.25)
    for r in p.runs:
        _normalize_body_run(r)


# ---------- tables ----------
def restyle_data_table(table) -> None:
    """Apply red header + zebra body to 3-col data tables (visitors, events)."""
    set_table_fixed_layout(table)
    # Header row.
    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALETTE["primary"])
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        set_cell_borders(
            cell,
            bottom={"sz": 8, "color": PALETTE["primary"], "val": "single"},
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, font="Cambria", size_pt=10, bold=True,
                          color=rgb(PALETTE["white"]), all_caps=True, tracking=20)
    set_row_height(table.rows[0], 360, exact=False)

    # Body rows: zebra (darker than masthead cream so stripes are visible).
    # Explicit shading on EVERY cell -- otherwise residual fills from the
    # original template (e.g. pale blue) survive on even rows.
    for ri, row in enumerate(table.rows[1:], start=1):
        zebra = (ri % 2 == 1)
        for cell in row.cells:
            set_cell_shading(
                cell, PALETTE["zebra"] if zebra else PALETTE["white"])
            set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
            set_cell_borders(
                cell,
                bottom={"sz": 4, "color": PALETTE["hairline"], "val": "single"},
            )
            for p in cell.paragraphs:
                for r in p.runs:
                    _normalize_body_run(r)
                    r.font.size = Pt(10)  # data tables run a touch smaller


def restyle_layout_table(table, *, label_color=False) -> None:
    """Apply borderless layout styling to 2-col layout tables.

    Clears any pre-existing cell shading from the original template
    (e.g. pale-blue or deep-blue Word default) so the new design ships
    consistently.
    """
    set_table_fixed_layout(table)
    remove_table_borders(table)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            set_cell_shading(cell, PALETTE["white"])
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            for p in cell.paragraphs:
                for r in p.runs:
                    _normalize_body_run(r)
                    if label_color and ci == 0 and r.bold:
                        r.font.color.rgb = PRIMARY


def _replace_paragraph_text(p, new_text: str) -> None:
    """Clear all runs in a paragraph and write `new_text` in its place,
    preserving the paragraph's existing style. Style is inherited from
    the first run if available."""
    # Capture style from the first run (if any) so we don't lose
    # bold / colour applied by the surrounding template restyle.
    sample = p.runs[0] if p.runs else None
    sample_font_name = sample.font.name if sample else None
    sample_size = sample.font.size if sample else None
    sample_bold = sample.bold if sample else None
    sample_italic = sample.italic if sample else None
    sample_color = (sample.font.color.rgb if sample and
                    sample.font.color and sample.font.color.rgb else None)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    new_run = p.add_run(new_text)
    if sample_font_name:
        new_run.font.name = sample_font_name
    if sample_size:
        new_run.font.size = sample_size
    if sample_bold is not None:
        new_run.bold = sample_bold
    if sample_italic is not None:
        new_run.italic = sample_italic
    if sample_color:
        new_run.font.color.rgb = sample_color


def insert_dean_name(table) -> None:
    """Replace [Dean's Name] / [Full Name] placeholders in Section 1.

    - Left cell (credentials block): "[Dean's Name]" -> DEAN_NAME
    - Right cell (message signature): "[Full Name], MD, PhD"
                                       -> "{DEAN_NAME_PLAIN}, MD, PhD"
    """
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                stripped = p.text.strip()
                if stripped == "[Dean's Name]":
                    _replace_paragraph_text(p, DEAN_NAME)
                elif stripped == "[Full Name], MD, PhD":
                    _replace_paragraph_text(
                        p, f"{DEAN_NAME_PLAIN}, MD, PhD")
                elif stripped.startswith("[Full Name]"):
                    _replace_paragraph_text(p, DEAN_NAME_PLAIN)


def insert_dean_photo(table) -> None:
    """Replace the '[ Photo ]' placeholder in Table 1 with the Dean image."""
    if not DEAN_PATH.exists():
        return
    cell = table.rows[0].cells[0]
    for p in cell.paragraphs:
        if p.text.strip() == DEAN_PHOTO_PLACEHOLDER:
            # Clear runs and insert image.
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=0, space_after=4)
            run = p.add_run()
            run.add_picture(str(DEAN_PATH), width=Inches(1.6))
            # Red frame around the photo cell as matte (per visual spec).
            set_cell_borders(
                cell,
                top={"sz": 8, "color": PALETTE["primary"], "val": "single"},
                bottom={"sz": 8, "color": PALETTE["primary"], "val": "single"},
                left={"sz": 8, "color": PALETTE["primary"], "val": "single"},
                right={"sz": 8, "color": PALETTE["primary"], "val": "single"},
            )
            return


def restyle_highlights_table(table) -> None:
    """Featured Highlights tables (3 cols: card, gutter, card)."""
    set_table_fixed_layout(table)
    remove_table_borders(table)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci == 1:
                continue  # gutter
            set_cell_shading(cell, PALETTE["cream"])
            set_cell_margins(cell, top=200, bottom=180, left=200, right=200)
            # Premium card affordance: gold rule on top instead of red bar
            # on left -- matches the HTML rendering and reads more editorial.
            set_cell_borders(
                cell,
                top={"sz": 18, "color": PALETTE["accent_aa"], "val": "single"},
            )
            for p in cell.paragraphs:
                for r in p.runs:
                    _normalize_body_run(r)


# ---------- header / footer ----------
def restyle_header_footer(doc) -> None:
    section = doc.sections[0]

    # HEADER
    header = section.header
    # Remove existing header paragraphs (keep one to add to).
    for p in list(header.paragraphs):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        # keep paragraph object
    p = header.paragraphs[0]
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("NAGOYA UNIVERSITY  ·  GRADUATE SCHOOL OF MEDICINE")
    style_run(r, font="Calibri", size_pt=8.5, color=MUTED,
              small_caps=True, tracking=40)
    set_paragraph_border(p, position="bottom", sz=4,
                         color=PALETTE["accent"], val="single", space=4)

    # FOOTER
    footer = section.footer
    for p in list(footer.paragraphs):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    p = footer.paragraphs[0]
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_paragraph_border(p, position="top", sz=4,
                         color=PALETTE["hairline"], val="single", space=4)
    parts_left = "Graduate School of Medicine, Nagoya University"
    parts_mid = "  ·  med.nagoya-u.ac.jp  ·  "
    r1 = p.add_run(parts_left)
    style_run(r1, font="Calibri", size_pt=8, color=MUTED)
    r2 = p.add_run(parts_mid)
    style_run(r2, font="Calibri", size_pt=8, color=MUTED)
    r3 = p.add_run("Page ")
    style_run(r3, font="Calibri", size_pt=8, color=MUTED)
    add_page_field(p, "PAGE")
    r4 = p.add_run(" of ")
    style_run(r4, font="Calibri", size_pt=8, color=MUTED)
    add_page_field(p, "NUMPAGES")
    # Apply Calibri 8pt to the field-generated runs (last 2 runs).
    for r in p.runs[-4:]:
        if r.font.size is None:
            r.font.size = Pt(8)
        if r.font.name in (None, ""):
            r.font.name = "Calibri"
        if r.font.color.rgb is None:
            r.font.color.rgb = MUTED


# ---------- page setup ----------
def configure_page(doc) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)


# ---------- main pipeline ----------
def build(src: Path = ORIGINAL_TEMPLATE, dst: Path = MERIDIAN_TEMPLATE) -> Path:
    if not src.exists():
        raise FileNotFoundError(f"Original template not found: {src}")

    # Work on a copy so we don't touch the original.
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    configure_page(doc)
    restyle_header_footer(doc)

    # 1) Masthead — Table 0.
    if doc.tables:
        restyle_masthead(doc.tables[TABLE_MASTHEAD])

    # 2) Body paragraphs — section heads, subheads, body, bullets.
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if is_section_heading(text):
            restyle_section_heading(p)
        elif text in SUBHEAD_TEXTS:
            restyle_subhead(p)
        elif p.style.name == "List Paragraph":
            restyle_bullet(p)
        else:
            restyle_body(p)

    # 3) Apply per-table styling. Named indices (from constants at the
    # top of this module) keep the orchestration grep-able.
    if TABLE_DEAN < len(doc.tables):
        restyle_layout_table(doc.tables[TABLE_DEAN], label_color=False)
        insert_dean_photo(doc.tables[TABLE_DEAN])
        insert_dean_name(doc.tables[TABLE_DEAN])
    for idx in (TABLE_HIGHLIGHTS_TOP, TABLE_HIGHLIGHTS_BOTTOM):
        if idx < len(doc.tables):
            restyle_highlights_table(doc.tables[idx])
    for idx in (TABLE_VISITORS, TABLE_EVENTS):
        if idx < len(doc.tables):
            restyle_data_table(doc.tables[idx])
    if TABLE_CONTACT < len(doc.tables):
        restyle_layout_table(doc.tables[TABLE_CONTACT], label_color=True)

    doc.save(str(dst))
    return dst


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    out = build()
    logging.getLogger(__name__).info("Built: %s", out)
