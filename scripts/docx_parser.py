"""Parse a filled newsletter DOCX into a structured Newsletter object.

Sections are keyed by their numbered heading (1..7). The original section
names and content are preserved verbatim.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from html import escape
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from scripts.config import SUBHEAD_TEXTS  # known sub-headings (canonical template)

log = logging.getLogger(__name__)


SECTION_HEAD_RE = re.compile(r"^\s*(\d+)\s*[\.\-—]?\s*[—-]?\s*(.+?)\s*$")
# `01 — TITLE` / `01 - TITLE` / `01. TITLE` / `01: TITLE` / fullwidth digits.
NUMBERED_HEAD_RE = re.compile(
    r"^\s*0?(\d+)\s*[—–\-:.]\s+(.+?)\s*$"
)
# Three alternative shapes for legacy / non-numbered-prefix section heads.
# The bare-numeric alternative REQUIRES an explicit separator -- without it
# `1 Recent grant from JSPS` (a body sentence beginning with a digit) would
# get mis-parsed as section 1 titled "Recent grant from JSPS".
#
#   1. English prose prefix: `Section 5: Title` / `Sec. 5 — Title`
#   2. Japanese: `第N章 Title` / `第N号 — Title` / `第N Title`
#      (the `第` prefix itself is the marker -- kanji suffix optional,
#      separator optional, since this form rarely uses ASCII punctuation)
#   3. Bare numeric: `1. Title` / `5: Title` / `7 — Title`
#      Separator is mandatory here.
LEGACY_HEAD_RE = re.compile(
    r"^\s*(?:"
    r"(?:Section|Sec\.?)\s+(?P<en_num>\d+)\s*[\.:—–\-]?\s*"
    r"(?P<en_title>.+?)"
    r"|"
    r"第\s*(?P<jp_num>\d+)\s*[章号節]?\s*[\.:—–\-]?\s*"
    r"(?P<jp_title>.+?)"
    r"|"
    r"(?P<num>\d+)\s*[\.:—–\-]\s+(?P<title>.+?)"
    r")\s*$"
)


# ---------- block dataclasses ----------
@dataclass(frozen=True)
class HtmlText:
    """A piece of inline HTML (already escaped + with run formatting tags)."""

    html: str


@dataclass(frozen=True)
class Heading:
    level: int
    text: str


@dataclass(frozen=True)
class BodyParagraph:
    html: str


@dataclass(frozen=True)
class BulletList:
    items: tuple[str, ...]  # each item is HTML-safe


@dataclass(frozen=True)
class TableBlock:
    rows: tuple[tuple[str, ...], ...]  # rows of HTML cells
    has_header: bool


@dataclass(frozen=True)
class ImageRef:
    rel_id: str          # docx relationship id
    filename: str        # basename inside word/media/
    alt: str = ""
    url: str = ""        # public URL (filled by renderer after image extraction)


Block = BodyParagraph | BulletList | TableBlock | ImageRef | Heading


@dataclass(frozen=True)
class Section:
    number: int
    title: str
    blocks: tuple[Block, ...]


@dataclass(frozen=True)
class Masthead:
    title: str
    tagline: str
    subtitle: str
    issue_line: str


@dataclass(frozen=True)
class Newsletter:
    masthead: Masthead
    sections: tuple[Section, ...]


# ---------- run → HTML ----------
def _run_to_html(run) -> str:
    text = escape(run.text or "")
    if not text:
        return ""
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    return text


# ---------- inline image (drawing → <img>) ----------
DRAWING_TAG = qn("w:drawing")
BLIP_TAG = qn("a:blip")
EMBED_ATTR = qn("r:embed")
EXTENT_TAG = qn("wp:extent")
PIC_CNVPR_TAG = qn("pic:cNvPr")

# EMU per pixel at 96 dpi (914400 EMU / inch / 96 px / inch).
EMU_PER_PX = 9525
# Cap image width to fit the 600 px email container with padding.
MAX_IMG_PX = 560


def _drawing_to_img(drawing, part) -> str:
    """Return an <img> tag (with media:// sentinel src) for a w:drawing."""
    blip = drawing.find(".//" + BLIP_TAG)
    if blip is None:
        return ""
    rid = blip.get(EMBED_ATTR)
    if not rid:
        return ""
    rel = part.rels.get(rid)
    if rel is None or not rel.target_ref:
        return ""
    fname = Path(rel.target_ref).name

    # Optional alt text from pic:cNvPr@descr or @name.
    alt = ""
    cnv_pr = drawing.find(".//" + PIC_CNVPR_TAG)
    if cnv_pr is not None:
        alt = cnv_pr.get("descr") or cnv_pr.get("name") or ""

    # Width and height from wp:extent (in EMU). Cap width to MAX_IMG_PX
    # and scale height proportionally so Outlook (which ignores
    # `height:auto`) reserves the right vertical space even when images
    # are blocked.
    size_attrs = ""
    extent = drawing.find(".//" + EXTENT_TAG)
    if extent is not None:
        cx_str, cy_str = extent.get("cx"), extent.get("cy")
        if cx_str and cx_str.isdigit():
            cx = int(cx_str)
            width_px = max(1, min(cx // EMU_PER_PX, MAX_IMG_PX))
            size_attrs = f' width="{width_px}"'
            if cy_str and cy_str.isdigit():
                cy = int(cy_str)
                # Scale height proportionally if width was capped.
                ratio = width_px / (cx / EMU_PER_PX)
                height_px = max(1, int(cy / EMU_PER_PX * ratio))
                size_attrs += f' height="{height_px}"'

    return (
        f'<img src="media://{escape(fname, quote=True)}" '
        f'alt="{escape(alt, quote=True)}"{size_attrs} '
        f'style="display:block;max-width:100%;height:auto;'
        f'margin:0;border:0;" />'
    )


def _hyperlinks(paragraph: Paragraph) -> dict[str, str]:
    """Map relationship ids to URLs for hyperlinks in this paragraph."""
    out = {}
    for hl in paragraph._p.findall(qn("w:hyperlink")):
        rid = hl.get(qn("r:id"))
        if not rid:
            continue
        rel = paragraph.part.rels.get(rid)
        if rel is not None:
            out[rid] = rel.target_ref
    return out


def paragraph_to_html(paragraph: Paragraph) -> str:
    """Convert a paragraph's runs (and hyperlinks) into safe HTML."""
    rid_to_url = _hyperlinks(paragraph)

    parts: list[str] = []
    for child in paragraph._p.iterchildren():
        tag = child.tag
        if tag == qn("w:r"):
            # Inline drawing inside this run? Emit an <img> tag.
            drawing = child.find(".//" + DRAWING_TAG)
            if drawing is not None:
                img = _drawing_to_img(drawing, paragraph.part)
                if img:
                    parts.append(img)
                    continue
            for r in paragraph.runs:
                if r._r is child:
                    parts.append(_run_to_html(r))
                    break
        elif tag == qn("w:hyperlink"):
            url = rid_to_url.get(child.get(qn("r:id")), "")
            inner_runs = []
            for r_el in child.findall(qn("w:r")):
                # Build text manually since Paragraph.runs doesn't include
                # runs inside hyperlinks.
                texts = [t.text or "" for t in r_el.findall(qn("w:t"))]
                inner_runs.append(escape("".join(texts)))
            label = "".join(inner_runs) or escape(url)
            if url:
                parts.append(f'<a href="{escape(url, quote=True)}">{label}</a>')
            else:
                parts.append(label)
    return "".join(parts).strip()


# ---------- helpers ----------
def _is_list_paragraph(p: Paragraph) -> bool:
    return (p.style.name or "").startswith("List Paragraph")


def _detect_section(text: str) -> tuple[int, str] | None:
    """Detect a section heading like '01 — RESEARCH' or '1. Research'.

    Tries the strict numeric pattern first; falls back to the legacy
    pattern that supports English `Section N` and Japanese `第N章`
    prefixes plus a separator-required bare-numeric form.
    """
    m = NUMBERED_HEAD_RE.match(text)
    if m:
        return int(m.group(1)), m.group(2).strip()
    m = LEGACY_HEAD_RE.match(text)
    if m:
        num = m.group("en_num") or m.group("jp_num") or m.group("num")
        title = (m.group("en_title") or m.group("jp_title")
                 or m.group("title") or "")
        if num and title.strip():
            return int(num), title.strip()
    return None


# Maximum text length for a paragraph to be treated as a sub-heading via
# the structural heuristic. Prevents long sentences from being mistaken
# for headings just because they're bold.
_SUBHEAD_MAX_CHARS = 80


def is_subheading_paragraph(p: Paragraph, text: str | None = None) -> bool:
    """Decide whether a paragraph is a sub-heading.

    Detection is purely structural so editors can add / rename / remove
    sub-sections in Word and the toolkit picks them up automatically:

      1. Word's built-in `Heading 2` / `Heading 3` / ... styles.
      2. Backwards-compat: text exactly matches one of `SUBHEAD_TEXTS`.
      3. Short, all-bold paragraph that doesn't end like a sentence.

    Section-level headings ("1. ..." / "01 — ...") are detected
    separately in `_detect_section` and short-circuit before we get here.

    `text` is optional -- if not supplied we derive it from `p.text.strip()`.
    Pass it when you already have it to save one strip() call.
    """
    if text is None:
        text = p.text.strip()
    if not text or len(text) > _SUBHEAD_MAX_CHARS:
        return False

    # 1. Word built-in heading style (Heading 2/3/... -- never Heading 1
    # which we reserve for section headings).
    # `style_id` is the locale-invariant Word identifier
    # ("Heading2", "Heading3"...). `style.name` is localized
    # ("見出し 2", "Überschrift 2"). Prefer style_id; fall back to name.
    style = p.style
    style_id = getattr(style, "style_id", None) or ""
    style_name = (getattr(style, "name", "") or "")
    if style_id.startswith("Heading") and style_id != "Heading1":
        return True
    if style_name.startswith("Heading") and not style_name.endswith(" 1"):
        return True

    # 2. Backwards-compat: hard-coded list of canonical sub-headings.
    if text in SUBHEAD_TEXTS:
        return True

    # 3. Heuristic: short bold paragraph without sentence punctuation.
    runs_with_text = [r for r in p.runs if r.text.strip()]
    if not runs_with_text:
        return False
    if not all(bool(r.bold) for r in runs_with_text):
        return False
    if text.endswith((".", "!", "?", "…", "...")):
        return False
    # Avoid catching a single bold word inside a normal paragraph -- a
    # sub-head is usually a complete short label.
    if len(text) < 3:
        return False
    return True


# ---------- table → block ----------
def _table_to_block(table: Table) -> TableBlock:
    rows_out: list[tuple[str, ...]] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_html_parts = []
            for p in cell.paragraphs:
                ph = paragraph_to_html(p)
                if ph:
                    cell_html_parts.append(ph)
            cells.append("<br>".join(cell_html_parts))
        rows_out.append(tuple(cells))
    # First row is header if all cells are short labels (heuristic: <= 30 chars
    # and bold dominant) — for safety we say it's a header.
    has_header = len(rows_out) >= 2
    return TableBlock(rows=tuple(rows_out), has_header=has_header)


# ---------- masthead extraction ----------
def _extract_masthead(doc: DocxDocument) -> Masthead:
    """Pull title/tagline/subtitle/issue line from the first table."""
    if not doc.tables:
        return Masthead("", "", "", "")
    cell = doc.tables[0].rows[0].cells[1]
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else ""
    tagline = paragraphs[1] if len(paragraphs) > 1 else ""
    subtitle = paragraphs[2] if len(paragraphs) > 2 else ""
    issue_line = paragraphs[3] if len(paragraphs) > 3 else ""
    return Masthead(title, tagline, subtitle, issue_line)


# ---------- main parse ----------
def _iter_body_blocks(doc: DocxDocument) -> Iterable[tuple[str, object]]:
    """Yield ('paragraph', Paragraph) or ('table', Table) in document order."""
    body = doc.element.body
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    p_idx = 0
    t_idx = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p") and p_idx < len(paragraphs):
            yield "paragraph", paragraphs[p_idx]
            p_idx += 1
        elif child.tag == qn("w:tbl") and t_idx < len(tables):
            yield "table", tables[t_idx]
            t_idx += 1


def parse(docx_path: Path) -> Newsletter:
    """Parse a filled DOCX into a Newsletter."""
    doc = Document(str(docx_path))
    masthead = _extract_masthead(doc)

    sections: list[Section] = []
    current_num: int | None = None
    current_title: str = ""
    current_blocks: list[Block] = []
    pending_bullets: list[str] = []
    table_index = 0

    def flush_bullets():
        if pending_bullets:
            current_blocks.append(BulletList(items=tuple(pending_bullets)))
            pending_bullets.clear()

    def flush_section():
        nonlocal current_num, current_title, current_blocks
        if current_num is not None:
            sections.append(Section(
                number=current_num,
                title=current_title,
                blocks=tuple(current_blocks),
            ))
        current_blocks = []

    for kind, item in _iter_body_blocks(doc):
        if kind == "paragraph":
            p: Paragraph = item
            text = p.text.strip()
            sec = _detect_section(text)
            if sec is not None:
                flush_bullets()
                flush_section()
                current_num, current_title = sec
                continue
            if current_num is None:
                continue  # pre-section content (masthead is in tables[0])

            # Bullet list?
            if _is_list_paragraph(p):
                html = paragraph_to_html(p)
                if html:
                    pending_bullets.append(html)
                continue
            flush_bullets()

            # Subhead? (Word style + heuristic + canonical-template list)
            if is_subheading_paragraph(p, text):
                current_blocks.append(Heading(level=3, text=text))
                continue

            # `paragraph_to_html` now embeds inline images directly via
            # media:// sentinel URLs, so a stand-alone image paragraph
            # renders as <img> inside the body paragraph.
            html = paragraph_to_html(p)
            if html:
                current_blocks.append(BodyParagraph(html=html))

        else:  # table
            t: Table = item
            table_index += 1
            if table_index == 1:
                continue  # masthead table — already extracted
            if current_num is None:
                continue
            flush_bullets()
            current_blocks.append(_table_to_block(t))

    flush_bullets()
    flush_section()

    # Section count is intentionally flexible -- log at DEBUG so editors
    # who add/remove sections don't see noisy warnings on every build.
    log.debug("Parsed %d section(s)", len(sections))

    return Newsletter(masthead=masthead, sections=tuple(sections))


__all__ = [
    "Newsletter", "Masthead", "Section", "Heading",
    "BodyParagraph", "BulletList", "TableBlock", "ImageRef",
    "parse", "paragraph_to_html",
    "is_subheading_paragraph",
]
