"""Build the modernized MERIDIAN newsletter DOCX template.

Reads the original NagoyaU template, preserves all section names and
content verbatim (text only), and restyles it per the Meridian visual
spec:

* Title MERIDIAN (replaces MEDICAL FRONTIER)
* NU blue (#003F88) + warm gold palette per the official guideline
* Cambria headings, Calibri body
* Masthead band, section bars, gold dividers
* Zebra-styled data tables, refined header/footer

Bundle 27 split this single 600-line module into a small package:

* `_styles.py`   -- color/run/paragraph primitives
* `_elements.py` -- per-element restylers (masthead, sections, tables)
* `__init__.py`  -- public API + the `build()` orchestration

The public surface is unchanged: callers still do
`from scripts import build_template as bt; bt.build(src, dst)`.
"""

from __future__ import annotations

import logging
import shutil
from pathlib import Path
from typing import Final

from docx import Document

from scripts.config import MERIDIAN_TEMPLATE, ORIGINAL_TEMPLATE
from scripts.docx_parser import is_subheading_paragraph

# Re-export the primitives + element restylers so existing test code
# / debug shells that hit `bt.<name>` keep working without changes.
from scripts.build_template._elements import (
    SECTION_HEAD_RE,
    configure_page,
    insert_dean_name,
    insert_dean_photo,
    is_section_heading,
    restyle_body,
    restyle_bullet,
    restyle_data_table,
    restyle_header_footer,
    restyle_highlights_table,
    restyle_layout_table,
    restyle_masthead,
    restyle_section_heading,
    restyle_subhead,
)
from scripts.build_template._styles import (
    ACCENT, MUTED, PRIMARY, TEXT,
    _normalize_body_run, rgb, style_paragraph, style_run,
)

log = logging.getLogger(__name__)


# Index of each table inside the DOCX (in document order). Named so the
# orchestration in `build()` is grep-able instead of mystery-meat
# `doc.tables[3]` calls. `Final` so type-checkers flag accidental
# reassignment.
TABLE_MASTHEAD: Final[int] = 0
TABLE_DEAN: Final[int] = 1
TABLE_HIGHLIGHTS_TOP: Final[int] = 2
TABLE_HIGHLIGHTS_BOTTOM: Final[int] = 3
TABLE_VISITORS: Final[int] = 4
TABLE_EVENTS: Final[int] = 5
TABLE_CONTACT: Final[int] = 6


def build(src: Path = ORIGINAL_TEMPLATE, dst: Path = MERIDIAN_TEMPLATE) -> Path:
    """Read the original DOCX, restyle it per the visual spec, save to `dst`.

    Returns the destination path. The source file is never modified --
    we work on a copy.
    """
    if not src.exists():
        raise FileNotFoundError(f"Original template not found: {src}")

    # Work on a copy so we don't touch the original.
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    configure_page(doc)
    restyle_header_footer(doc)

    # 1) Masthead — Table 0.
    if doc.tables:
        restyle_masthead(doc.tables[TABLE_MASTHEAD])

    # 2) Body paragraphs -- section heads, subheads, body, bullets.
    # Sub-heading detection delegates to docx_parser.is_subheading_paragraph
    # (same 3-tier check the parser uses) so the styled template stays in
    # sync with what the parser will recognise as sub-headings in any
    # issue's filled-in DOCX.
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if is_section_heading(text):
            restyle_section_heading(p)
        elif is_subheading_paragraph(p, text):
            restyle_subhead(p)
        elif p.style.name == "List Paragraph":
            restyle_bullet(p)
        else:
            restyle_body(p)

    # 3) Apply per-table styling. Named indices (from constants at the
    # top of this module) keep the orchestration grep-able.
    if TABLE_DEAN < len(doc.tables):
        restyle_layout_table(doc.tables[TABLE_DEAN], label_color=False)
        insert_dean_photo(doc.tables[TABLE_DEAN])
        insert_dean_name(doc.tables[TABLE_DEAN])
    for idx in (TABLE_HIGHLIGHTS_TOP, TABLE_HIGHLIGHTS_BOTTOM):
        if idx < len(doc.tables):
            restyle_highlights_table(doc.tables[idx])
    for idx in (TABLE_VISITORS, TABLE_EVENTS):
        if idx < len(doc.tables):
            restyle_data_table(doc.tables[idx])
    if TABLE_CONTACT < len(doc.tables):
        restyle_layout_table(doc.tables[TABLE_CONTACT], label_color=True)

    doc.save(str(dst))
    return dst


__all__ = [
    "build",
    "TABLE_MASTHEAD", "TABLE_DEAN",
    "TABLE_HIGHLIGHTS_TOP", "TABLE_HIGHLIGHTS_BOTTOM",
    "TABLE_VISITORS", "TABLE_EVENTS", "TABLE_CONTACT",
    # Re-exported helpers (kept for debug shells / future test reuse).
    "rgb", "PRIMARY", "ACCENT", "TEXT", "MUTED",
    "style_run", "style_paragraph", "_normalize_body_run",
    "is_section_heading", "SECTION_HEAD_RE",
    "restyle_masthead", "restyle_section_heading", "restyle_subhead",
    "restyle_body", "restyle_bullet",
    "restyle_data_table", "restyle_layout_table",
    "restyle_highlights_table", "restyle_header_footer",
    "insert_dean_name", "insert_dean_photo",
    "configure_page",
]


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    out = build()
    log.info("Built: %s", out)
