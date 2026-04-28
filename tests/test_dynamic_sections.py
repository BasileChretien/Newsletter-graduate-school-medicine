"""Tests for editor-defined sections.

The toolkit supports any number of sections, any titles, any
sub-headings -- as long as the editor uses Word's heading styles or
keeps the "01 -- TITLE" / "1. Title" numbered-heading pattern. These
tests cover the flexibility the rebrand introduced.
"""

from __future__ import annotations

from pathlib import Path

import docx  # python-docx is a hard requirement, not optional

from scripts.docx_parser import Heading, parse


def _make_doc(tmp_path: Path,
              paragraphs: list[tuple[str, str | None]]) -> Path:
    """Build a minimal one-table-masthead DOCX from a list of
    (text, style_name) tuples. style_name=None -> Normal paragraph."""
    doc = docx.Document()
    # Masthead table (the parser skips table_index == 1).
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[1].text = "MERIDIAN\nTagline\nSubtitle\nVOL. X"
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style is not None:
            try:
                p.style = doc.styles[style]
            except KeyError:
                pass
    out = tmp_path / "synthetic.docx"
    doc.save(str(out))
    return out


def test_three_sections_with_arbitrary_titles(tmp_path: Path) -> None:
    """Editor adds three custom sections; toolkit accepts all three."""
    doc_path = _make_doc(tmp_path, [
        ("1.  Hospital News", None),
        ("Some hospital body text.", None),
        ("2.  Awards & Honours", None),
        ("Some awards body text.", None),
        ("3.  Save the Date", None),
        ("Some events body text.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 3
    assert nl.sections[0].title == "Hospital News"
    assert nl.sections[1].title == "Awards & Honours"
    assert nl.sections[2].title == "Save the Date"


def test_subheading_via_word_heading_style(tmp_path: Path) -> None:
    """Word's built-in 'Heading 2' style is recognised as a sub-heading."""
    doc_path = _make_doc(tmp_path, [
        ("1.  Free Section", None),
        ("Sub-heading via Word style", "Heading 2"),
        ("Body paragraph below the sub-heading.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 1
    headings = [b.text for b in nl.sections[0].blocks
                if isinstance(b, Heading)]
    assert "Sub-heading via Word style" in headings


def test_subheading_via_bold_short_paragraph(tmp_path: Path) -> None:
    """Heuristic: short bold paragraph without sentence punctuation."""
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[1].text = "MERIDIAN\nTagline\nSubtitle\nVOL. X"
    doc.add_paragraph("1.  Free Section")
    p = doc.add_paragraph()
    run = p.add_run("Bold Subhead")
    run.bold = True
    doc.add_paragraph("Body paragraph below.")
    out = tmp_path / "syn.docx"
    doc.save(str(out))

    nl = parse(out)
    headings = [b.text for b in nl.sections[0].blocks
                if isinstance(b, Heading)]
    assert "Bold Subhead" in headings


def test_long_bold_paragraph_is_NOT_a_subhead(tmp_path: Path) -> None:
    """A long bold paragraph (e.g. an emphasised body sentence) is body, not heading."""
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[1].text = "MERIDIAN\nTagline\nSubtitle\nVOL. X"
    doc.add_paragraph("1.  Section")
    p = doc.add_paragraph()
    long_text = (
        "This sentence is bold for emphasis but it is much longer than "
        "any reasonable sub-heading and ends with a period."
    )
    run = p.add_run(long_text)
    run.bold = True
    out = tmp_path / "syn.docx"
    doc.save(str(out))

    nl = parse(out)
    headings = [b.text for b in nl.sections[0].blocks
                if isinstance(b, Heading)]
    assert long_text not in headings


def test_section_heading_with_colon(tmp_path: Path) -> None:
    """`8: Title` (colon, common Japanese habit) is recognised as a section."""
    doc_path = _make_doc(tmp_path, [
        ("1: Opening Section", None),
        ("Body of section one.", None),
        ("2: Closing Section", None),
        ("Body of section two.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 2
    assert [s.title for s in nl.sections] == ["Opening Section", "Closing Section"]


def test_section_heading_with_section_prefix(tmp_path: Path) -> None:
    """`Section N. Title` (English-style prose prefix) is recognised."""
    doc_path = _make_doc(tmp_path, [
        ("Section 1. Opening Section", None),
        ("Body.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 1
    assert nl.sections[0].title == "Opening Section"


def test_section_heading_japanese_dai_n_shou(tmp_path: Path) -> None:
    """`第N章 Title` (Japanese chapter prefix) is recognised."""
    doc_path = _make_doc(tmp_path, [
        ("第3章 Lab News", None),
        ("Body of the Japanese-titled section.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 1
    assert nl.sections[0].number == 3
    assert nl.sections[0].title == "Lab News"


def test_section_heading_japanese_dai_n_gou_with_dash(tmp_path: Path) -> None:
    """`第N号 — Title` (Japanese issue prefix + em-dash) is recognised."""
    doc_path = _make_doc(tmp_path, [
        ("第8号 — タイトル", None),
        ("Body.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 1
    assert nl.sections[0].number == 8
    assert nl.sections[0].title == "タイトル"


def test_section_heading_japanese_no_kanji_suffix(tmp_path: Path) -> None:
    """`第3 Title` (no kanji suffix, plain digit + title) still works."""
    doc_path = _make_doc(tmp_path, [
        ("第3 Lab News", None),
        ("Body.", None),
    ])
    nl = parse(doc_path)
    assert len(nl.sections) == 1
    assert nl.sections[0].title == "Lab News"
