"""Integration-flavoured tests for `build_newsletter._build_pipeline`.

The pipeline is short but the ordering of validate -> write matters:
when validation hard-blocks (e.g. unfilled masthead `VOL. XX`), the
HTML must NOT be persisted to disk -- otherwise the editor's
double-click on `dist/issue-N.html` opens a stale broken file.
"""

from __future__ import annotations

from pathlib import Path

import docx

import build_newsletter as bn


def _set_masthead_lines(cell, lines: list[str]) -> None:
    """python-docx's `cell.text = "a\\nb"` makes ONE paragraph with line
    breaks; the parser expects four separate paragraphs in the right
    masthead cell. Add them explicitly."""
    # Clear the auto-created empty paragraph and add real ones.
    cell.text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)


def _make_unfilled_docx(tmp_path: Path) -> Path:
    """Synthesize a DOCX with the masthead placeholders still unfilled."""
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=2)
    # Masthead cell: VOL. XX still present -> validate() must hard-block.
    _set_masthead_lines(t.rows[0].cells[1], [
        "MERIDIAN",
        "Where medicine meets the world.",
        "Newsletter of the Graduate School of Medicine",
        "VOL. XX | ISSUE NO. XX | MONTH YEAR",
    ])
    doc.add_paragraph("1.  Some Section")
    doc.add_paragraph("Body content for the section.")
    out = tmp_path / "unfilled.docx"
    doc.save(str(out))
    return out


def _make_filled_docx(tmp_path: Path) -> Path:
    """Synthesize a DOCX with a real masthead -- validation should pass."""
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=2)
    _set_masthead_lines(t.rows[0].cells[1], [
        "MERIDIAN",
        "Where medicine meets the world.",
        "Newsletter of the Graduate School of Medicine",
        "VOL. 12 | ISSUE NO. 3 | MARCH 2026",
    ])
    doc.add_paragraph("1.  Some Section")
    doc.add_paragraph("Body content for the section.")
    out = tmp_path / "filled.docx"
    doc.save(str(out))
    return out


def test_build_pipeline_does_not_write_when_validation_fails(
    tmp_path: Path, monkeypatch
) -> None:
    """Validation hard-block must NOT leave a stale HTML on disk.

    Regression guard for round-7 review finding: previously the
    pipeline wrote `dist/issue-N.html` first, then validated -- so an
    unfilled-masthead build still left an openable file that the
    editor could double-click and paste into Outlook.
    """
    monkeypatch.setattr(bn, "DIST_DIR", tmp_path / "dist")
    monkeypatch.setattr(bn, "ASSETS_DIR", tmp_path / "assets")
    monkeypatch.setattr(bn, "DROP_DIR", tmp_path / "drop-images")

    docx_path = _make_unfilled_docx(tmp_path)
    result = bn._build_pipeline(docx_path, issue=1, validate_remote=False)

    assert result.exit_code != 0, "unfilled masthead must fail the build"
    out_html = tmp_path / "dist" / "issue-1.html"
    assert not out_html.exists(), (
        "Validation failure should not leave a stale HTML file on disk"
    )


def test_build_pipeline_writes_when_validation_passes(
    tmp_path: Path, monkeypatch
) -> None:
    """Sanity check: a DOCX with a filled masthead does produce a file."""
    monkeypatch.setattr(bn, "DIST_DIR", tmp_path / "dist")
    monkeypatch.setattr(bn, "ASSETS_DIR", tmp_path / "assets")
    monkeypatch.setattr(bn, "DROP_DIR", tmp_path / "drop-images")

    docx_path = _make_filled_docx(tmp_path)
    result = bn._build_pipeline(docx_path, issue=2, validate_remote=False)

    assert result.exit_code == 0
    out_html = tmp_path / "dist" / "issue-2.html"
    assert out_html.exists()
    assert "MERIDIAN" in out_html.read_text(encoding="utf-8")
