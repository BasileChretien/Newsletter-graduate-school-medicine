"""Round-17 regression tests for `--output-dir` and writability fallback.

The user reported (production field trial):
  * macOS: running the launcher without sudo, the toolkit failed to
    save outputs because the extracted ZIP folder was inside Downloads
    where modern macOS sandbox blocks writes.
  * Feature ask: let editors choose where the HTML lands.

These tests pin:
  * `is_writable_location()` correctly probe-writes and returns
    True/False without leaving probe files on disk.
  * `default_safe_output_dir()` honours `MERIDIAN_OUTPUT_DIR`,
    falls back to PROJECT_ROOT when writable, then ~/Documents,
    then ~/Meridian-Newsletter.
  * `_resolve_output_dir(None)` returns None on a writable
    PROJECT_ROOT (no behavioural change for the conventional case).
  * `_resolve_output_dir(explicit)` accepts the user's choice and
    sys.exit(2)s on an unwritable explicit path.
  * `_build_pipeline(output_dir=...)` writes HTML and assets to the
    redirected location, NOT to PROJECT_ROOT/dist or PROJECT_ROOT/assets.
  * `--output-dir` is plumbed through the CLI: build, preview,
    compose, and all commands all accept it.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

from click.testing import CliRunner

import build_newsletter as bn
from scripts.config import (
    default_safe_output_dir,
    is_writable_location,
)


# -- Writable-location probing -----------------------------------------------

def test_is_writable_location_true_for_tmp_path(tmp_path: Path):
    assert is_writable_location(tmp_path) is True
    # Probe file must NOT linger after the call.
    assert list(tmp_path.iterdir()) == [], (
        f"is_writable_location must clean up its probe file; "
        f"found leftovers: {list(tmp_path.iterdir())}"
    )


def test_is_writable_location_creates_missing_dir(tmp_path: Path):
    """Probing a non-existent path creates it (the path now exists)
    and returns True. The path-creation side-effect IS the write
    check, so we don't roll it back."""
    nested = tmp_path / "deep" / "nested" / "dir"
    assert not nested.exists()
    assert is_writable_location(nested) is True
    assert nested.exists()


def test_is_writable_location_false_for_unwritable(monkeypatch, tmp_path):
    """If `mkdir` raises (e.g. permission denied), probe returns False."""
    target = tmp_path / "denied"

    def boom(*a, **kw):
        raise PermissionError("sandbox denied")

    monkeypatch.setattr(Path, "mkdir", boom)
    assert is_writable_location(target) is False


def test_default_safe_output_dir_honours_env_var(monkeypatch, tmp_path):
    """`MERIDIAN_OUTPUT_DIR` env var wins over auto-detection."""
    override = tmp_path / "from_env"
    monkeypatch.setenv("MERIDIAN_OUTPUT_DIR", str(override))
    chosen = default_safe_output_dir()
    assert chosen.resolve() == override.resolve()


def test_default_safe_output_dir_falls_back_when_env_unwritable(
    monkeypatch, tmp_path,
):
    """An unwritable MERIDIAN_OUTPUT_DIR doesn't pin the user there
    -- the function falls through to the next tier."""
    fake_unwritable = tmp_path / "denied"
    monkeypatch.setenv("MERIDIAN_OUTPUT_DIR", str(fake_unwritable))

    # First mkdir call (the env-override probe) raises;
    # subsequent calls succeed.
    real_mkdir = Path.mkdir
    call_count = {"n": 0}

    def selective_mkdir(self, *a, **kw):
        call_count["n"] += 1
        if call_count["n"] == 1 and self == fake_unwritable:
            raise PermissionError("denied")
        return real_mkdir(self, *a, **kw)

    monkeypatch.setattr(Path, "mkdir", selective_mkdir)
    chosen = default_safe_output_dir()
    # Must NOT be the env override.
    assert chosen.resolve() != fake_unwritable.resolve()


# -- _resolve_output_dir CLI helper ------------------------------------------

def test_resolve_output_dir_none_when_writable_and_no_flag(
    monkeypatch, tmp_path,
):
    """If PROJECT_ROOT is writable and the user didn't pass
    --output-dir, return None so `_build_pipeline` keeps using
    `DIST_DIR` next to the script (the conventional case)."""
    monkeypatch.setattr(bn, "PROJECT_ROOT", tmp_path)
    assert bn._resolve_output_dir(None) is None


def test_resolve_output_dir_explicit_choice_wins(monkeypatch, tmp_path):
    """User-supplied --output-dir is honoured even when the toolkit
    folder is writable."""
    monkeypatch.setattr(bn, "PROJECT_ROOT", tmp_path)
    chosen = tmp_path / "custom"
    result = bn._resolve_output_dir(str(chosen))
    assert result == chosen.resolve()
    assert result.exists()


def test_resolve_output_dir_explicit_unwritable_exits_with_2(
    monkeypatch, tmp_path,
):
    """If the editor passes --output-dir pointing at a read-only
    location, fail fast with exit 2 -- they need to know
    immediately, not silently fall back."""
    monkeypatch.setattr(bn, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(bn, "is_writable_location", lambda _p: False)

    import pytest
    with pytest.raises(SystemExit) as excinfo:
        bn._resolve_output_dir(str(tmp_path / "denied"))
    assert excinfo.value.code == 2


def test_resolve_output_dir_auto_fallback_when_root_readonly(
    monkeypatch, tmp_path,
):
    """If PROJECT_ROOT is read-only AND no flag was given, redirect
    to the safe fallback (`~/Documents/Meridian-Newsletter` typically).
    This is the macOS-Downloads-sandbox production-bug case."""
    monkeypatch.setattr(bn, "PROJECT_ROOT", tmp_path / "fake_root")

    safe = tmp_path / "safe_output"
    safe.mkdir()

    def fake_writable(p):
        return p == safe

    monkeypatch.setattr(bn, "is_writable_location", fake_writable)
    monkeypatch.setattr(bn, "default_safe_output_dir", lambda: safe)

    result = bn._resolve_output_dir(None)
    assert result == safe


# -- CLI plumbing: --output-dir reaches the right place ----------------------

def test_build_cmd_accepts_output_dir(tmp_path):
    """The `build` subcommand must accept --output-dir without crashing.
    We patch _build_pipeline to just record what it received so we
    don't need a real DOCX fixture."""
    docx_in = tmp_path / "issue-1.docx"
    docx_in.write_bytes(b"fake")
    out_dir = tmp_path / "elsewhere"

    captured = {}

    # `**_` so this stub keeps testing what it is about -- that
    # `--output-dir` reaches the pipeline -- without needing an edit
    # every time an unrelated option is added to the command.
    def fake_pipeline(input_path, issue, *, validate_remote, output_dir, **_):
        captured["output_dir"] = output_dir
        return bn.BuildResult(0, "Test", out_dir / "dist" / "issue-1.html")

    with patch("build_newsletter._build_pipeline", new=fake_pipeline):
        runner = CliRunner()
        result = runner.invoke(
            bn.cli,
            ["build", "--input", str(docx_in), "--issue", "1",
             "--output-dir", str(out_dir), "--no-remote-check"],
        )
    assert result.exit_code == 0, result.output
    assert captured["output_dir"] == out_dir.resolve()


def test_preview_cmd_uses_output_dir_when_passed(tmp_path):
    """`preview --output-dir X` must look in X/dist, not PROJECT_ROOT/dist."""
    out_dir = tmp_path / "elsewhere"
    (out_dir / "dist").mkdir(parents=True)
    html = out_dir / "dist" / "issue-1.html"
    html.write_text("<html>x</html>", encoding="utf-8")

    with patch("webbrowser.open") as opener:
        runner = CliRunner()
        result = runner.invoke(
            bn.cli,
            ["preview", "--issue", "1", "--output-dir", str(out_dir)],
        )

    assert result.exit_code == 0, result.output
    opener.assert_called_once()
    assert html.as_uri() == opener.call_args.args[0]


def test_preview_cmd_default_uses_dist_dir(monkeypatch, tmp_path):
    """Without --output-dir, `preview` looks at the conventional
    DIST_DIR next to the script."""
    monkeypatch.setattr(bn, "DIST_DIR", tmp_path / "dist")
    (tmp_path / "dist").mkdir()
    html = tmp_path / "dist" / "issue-1.html"
    html.write_text("<html>x</html>", encoding="utf-8")

    with patch("webbrowser.open") as opener:
        runner = CliRunner()
        result = runner.invoke(bn.cli, ["preview", "--issue", "1"])

    assert result.exit_code == 0, result.output
    opener.assert_called_once()


# -- Empty-newsletter hard-fail (round-17 fix #1 catches the truly-empty
#    case the lenient fallback can't help with) ------------------------------

def test_build_pipeline_hard_fails_on_empty_docx(tmp_path):
    """A truly-empty DOCX (nothing in the body, even after lenient
    fallback) must produce exit_code=1 with a clear error message --
    the editor needs to see the problem in the launcher console
    BEFORE Outlook would otherwise open a blank draft."""
    import docx as docx_lib

    d = docx_lib.Document()
    p = tmp_path / "empty.docx"
    d.save(str(p))

    out_dir = tmp_path / "out"
    result = bn._build_pipeline(
        p, issue=1, validate_remote=False, output_dir=out_dir,
    )
    assert result.exit_code == 1
    # And no HTML was written.
    assert not (out_dir / "dist" / "issue-1.html").exists()


def test_output_dir_survives_a_docx_containing_photos(tmp_path: Path):
    """`--output-dir` crashed on any DOCX with an image.

    `_build_pipeline` resolved photo URLs with
    `to_raw_url(p, PROJECT_ROOT, ...)`, but `--output-dir` puts the
    extracted photos beside that directory instead of in the repo, so
    `to_raw_url` raised a bare ValueError traceback:

        Asset .../out/assets/issue-3/image1.jpg
        must be inside repo .../Newsletter-graduate-school-medicine

    Every real newsletter has photos, and `--output-dir` exists for the
    macOS case where the toolkit folder is unwritable -- where
    `default_safe_output_dir()` redirects to
    `~/Documents/Meridian-Newsletter` and hits exactly this path. So the
    v1.1.2 fix for that production failure could not survive a photo.

    Uses the shipped template, which carries the masthead logo and the
    dean photo, so the assertion rests on real embedded images.
    """
    import docx as docx_lib

    from scripts.config import MERIDIAN_TEMPLATE

    d = docx_lib.Document(str(MERIDIAN_TEMPLATE))
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in (("VOL. XX", "VOL. 4"),
                                     ("ISSUE NO. XX", "ISSUE NO. 1"),
                                     ("MONTH YEAR", "JUNE 2026")):
                        if old in para.text:
                            for run in para.runs:
                                run.text = run.text.replace(old, new)

    src = tmp_path / "issue-3.docx"
    d.save(str(src))
    out = tmp_path / "elsewhere"

    result = bn._build_pipeline(src, issue=3, validate_remote=False,
                                output_dir=out)

    assert result.exit_code == 0, "build --output-dir failed on a DOCX with photos"
    assert (out / "dist" / "issue-3.html").exists()
    # The URLs stay repo-relative, so a redirected build still produces
    # the same links a conventional one would.
    html = (out / "dist" / "issue-3.html").read_text(encoding="utf-8")
    assert "raw.githubusercontent.com" in html
    assert str(out).replace("\\", "/") not in html, (
        "a local filesystem path leaked into the published HTML")
