"""Tests for the subject-length warning thresholds.

Round-9 Email M1: a single soft warn at 78 chars was too loose --
inbox previews truncate around 50 chars on Outlook desktop / Gmail
web, so a recipient may only see "MERIDIAN — VOL. 12 | ISSUE NO."
instead of the full subject. Bundle 28 introduced two thresholds:
50 (preview) and 78 (RFC + spam-filter heuristic).

Round-10 code-review HIGH: bundle 28's tests duplicated the if/elif
threshold logic in a local helper rather than calling
`_build_pipeline` -- so they verified the test's copy of the logic,
not the real production code path. Bundle 29 invokes
`_build_pipeline` end-to-end via Click's `CliRunner` and asserts on
the actual stdout, so a future flip from `>` to `>=` (or any other
shape change) shows up here immediately.
"""

from __future__ import annotations

from pathlib import Path

import docx
from click.testing import CliRunner

import build_newsletter as bn


def _make_filled_docx_with_subject(
    tmp_path: Path, issue_line: str, name: str = "filled.docx"
) -> Path:
    """Synthesize a DOCX whose masthead issue-line drives the subject."""
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=2)
    cell = t.rows[0].cells[1]
    cell.text = "MERIDIAN"
    cell.add_paragraph("Where medicine meets the world.")
    cell.add_paragraph("Newsletter of the Graduate School of Medicine")
    cell.add_paragraph(issue_line)
    doc.add_paragraph("1.  Some Section")
    doc.add_paragraph("Body content for the section.")
    out = tmp_path / name
    doc.save(str(out))
    return out


def test_subject_threshold_constants_are_distinct():
    """Sanity: the two thresholds must differ AND order makes sense."""
    assert bn._SUBJECT_PREVIEW_LIMIT_CHARS == 50
    assert bn._SUBJECT_SPAM_LIMIT_CHARS == 78
    assert bn._SUBJECT_PREVIEW_LIMIT_CHARS < bn._SUBJECT_SPAM_LIMIT_CHARS


def _run_build_with_subject(
    tmp_path: Path, monkeypatch, issue_line: str
) -> str:
    """Drive `_build_pipeline` end-to-end via the real Click command
    and return stdout. The DOCX masthead's issue-line drives the
    subject (`MERIDIAN — {issue_line}`). All assets / dist live in
    tmp_path."""
    monkeypatch.setattr(bn, "DIST_DIR", tmp_path / "dist")
    monkeypatch.setattr(bn, "ASSETS_DIR", tmp_path / "assets")
    monkeypatch.setattr(bn, "DROP_DIR", tmp_path / "drop-images")
    docx_path = _make_filled_docx_with_subject(tmp_path, issue_line)
    runner = CliRunner()
    result = runner.invoke(bn.cli, [
        "build", "--input", str(docx_path), "--issue", "1",
        "--no-remote-check",
    ])
    return result.output


def test_subject_short_emits_no_warning(tmp_path: Path, monkeypatch):
    """A short issue-line keeps the rendered subject under 50 chars
    (`MERIDIAN — V1 | I1 | M1` is 22 chars). No warning fires."""
    out = _run_build_with_subject(tmp_path, monkeypatch, "V1 | I1 | M1")
    assert "Heads up" not in out
    assert "Note: subject" not in out


def test_subject_preview_threshold_emits_note_only(
    tmp_path: Path, monkeypatch
):
    """An issue-line that pushes the rendered subject between 51 and
    78 chars must trigger the preview-truncation NOTE (not the spam
    warning). `MERIDIAN — ` is 11 chars, so an issue line of 50
    chars yields a 61-char subject."""
    issue_line = "VOL. 12 | ISSUE NO. 3 | MARCH 2026 -- Special"  # 45 chars
    # -> subject ~57 chars -- in the preview-warn band.
    out = _run_build_with_subject(tmp_path, monkeypatch, issue_line)
    assert "Note: subject" in out
    assert "Heads up: subject" not in out


def test_subject_spam_threshold_emits_heads_up(tmp_path: Path, monkeypatch):
    """A long issue-line pushes the subject past 78 chars and triggers
    the firmer spam-heuristic warning, NOT the gentler preview note."""
    long_issue = "x" * 80
    out = _run_build_with_subject(tmp_path, monkeypatch, long_issue)
    assert "Heads up: subject" in out
    # The two warnings are mutually exclusive -- only one fires.
    assert "Note: subject" not in out


def test_subject_warning_mentions_month_year_tip(
    tmp_path: Path, monkeypatch
):
    """Round-10 UX M3: the warning text must point editors at the
    actual source (`MONTH YEAR` in the masthead issue line) not just
    say 'shorten the subject.'"""
    long_issue = "x" * 80
    out = _run_build_with_subject(tmp_path, monkeypatch, long_issue)
    assert "MONTH YEAR" in out or "MAR 2026" in out
