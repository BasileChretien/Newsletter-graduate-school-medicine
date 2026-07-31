"""The viewable copy of the newsletter, and the line between it and the email.

Background
----------
The renderer emits `<img src="https://raw.githubusercontent.com/...">`.
Those URLs only resolve once `publish-images` has pushed the photos --
and in CID mode nothing ever does, because the photos travel inside the
message instead. So the HTML on disk points at files that were never
uploaded.

The browser build hit this first, reported from the field as "the
preview is working perfectly, but when downloading the html the images
are gone". The desktop `preview` command had the same defect: it opened
`dist/issue-N.html`, which is that same document.

The fix is a second file, not a change to the first one -- and that
distinction is the thing most at risk of being "simplified" later, so it
is pinned hard below. `compose` reads `dist/issue-N.html` to build the
message, and CID mode works by rewriting exactly those URLs into `cid:`
references. If that file carried `data:` URIs instead, Outlook and Gmail
would strip them and every recipient would get a newsletter with no
photos at all -- a far worse bug than the one being fixed.
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest

import build_newsletter as bn
from scripts.mail.cid import InlineImage
from scripts.standalone import data_uri, to_standalone_html


def _img_srcs(html: str) -> list[str]:
    return re.findall(r'<img[^>]*\ssrc="([^"]*)"', html)


# ---------- the pure helper -------------------------------------------

def test_photo_urls_become_data_uris(tmp_path: Path):
    photo = tmp_path / "photo.jpg"
    photo.write_bytes(b"\xff\xd8\xff\xe0" + b"padding")
    url = "https://raw.githubusercontent.com/o/r/main/assets/issue-1/photo.jpg"
    html = f'<img src="{url}"><p>text</p>'

    out = to_standalone_html(
        html, [InlineImage(path=photo, cid="c1", original_url=url)])

    assert url not in out
    assert out.startswith('<img src="data:image/jpeg;base64,')
    assert "<p>text</p>" in out


def test_an_unreadable_photo_degrades_to_one_broken_image(tmp_path: Path):
    """Losing a photo must not lose the document. The email is built
    from a different file and is unaffected either way."""
    missing = tmp_path / "gone.jpg"
    url = "https://raw.githubusercontent.com/o/r/main/assets/issue-1/gone.jpg"
    html = f'<img src="{url}"><p>survives</p>'

    out = to_standalone_html(
        html, [InlineImage(path=missing, cid="c1", original_url=url)])

    assert "<p>survives</p>" in out
    assert url in out, "the original URL is left in place, not blanked"


def test_data_uri_carries_the_right_mime_type(tmp_path: Path):
    png = tmp_path / "logo.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n")
    assert data_uri(png).startswith("data:image/png;base64,")


# ---------- the build writes both, and they differ --------------------

@pytest.fixture(scope="module")
def built(tmp_path_factory):
    """One real build of the shipped template, reused by the tests below."""
    import docx as docx_lib

    from scripts.config import MERIDIAN_TEMPLATE

    d = docx_lib.Document(str(MERIDIAN_TEMPLATE))
    for para in d.paragraphs:
        for old, new in (("VOL. XX", "VOL. 2"),
                         ("ISSUE NO. XX", "ISSUE NO. 3"),
                         ("MONTH YEAR", "MARCH 2026")):
            if old in para.text:
                for run in para.runs:
                    run.text = run.text.replace(old, new)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in (("VOL. XX", "VOL. 2"),
                                     ("ISSUE NO. XX", "ISSUE NO. 3"),
                                     ("MONTH YEAR", "MARCH 2026")):
                        if old in para.text:
                            for run in para.runs:
                                run.text = run.text.replace(old, new)

    root = tmp_path_factory.mktemp("build")
    src = root / "issue-91.docx"
    d.save(str(src))
    # `output_dir` rather than the repo's own dist/: these tests must not
    # be able to overwrite or delete an editor's real issue-91, and the
    # fixture would otherwise leave extracted photos in the shared
    # assets/ directory.
    result = bn._build_pipeline(src, issue=91, validate_remote=False,
                                output_dir=root)
    yield (result,
           root / "dist" / "issue-91.html",
           root / "dist" / "issue-91.preview.html")


def test_the_build_writes_both_files(built):
    result, mail, standalone = built
    assert result.exit_code == 0
    assert mail.exists(), "the mail artefact must still be written"
    assert standalone.exists(), (
        "the viewable copy is missing -- `preview` falls back to the mail "
        "artefact, whose photo URLs were never published")


def test_the_mail_artefact_keeps_its_urls(built):
    """THE LOAD-BEARING TEST. `compose` reads this file and CID mode
    rewrites these URLs into `cid:` references. `data:` URIs here would
    be stripped by Outlook and Gmail, and every recipient would get a
    newsletter with no photos."""
    _, mail, _ = built
    html = mail.read_text(encoding="utf-8")

    assert "data:image/" not in html, (
        "the mail artefact must never carry data: URIs -- mail clients "
        "strip them, so recipients would lose every photo")
    assert any(s.startswith("https://raw.githubusercontent.com/")
               for s in _img_srcs(html))


def test_the_viewable_copy_needs_no_network(built):
    _, _, standalone = built
    html = standalone.read_text(encoding="utf-8")
    srcs = _img_srcs(html)

    assert srcs, "no <img> tags at all -- the build changed shape"
    assert all(s.startswith("data:image/") for s in srcs), (
        f"still fetching: {[s for s in srcs if not s.startswith('data:')]}")
    assert "raw.githubusercontent.com" not in html


def test_the_two_files_differ_only_in_their_image_sources(built):
    """Guards against the viewable copy drifting into a second render --
    it must be the same document, photos aside."""
    _, mail, standalone = built
    strip = lambda h: re.sub(r'\ssrc="[^"]*"', "", h)  # noqa: E731
    assert strip(mail.read_text(encoding="utf-8")) == \
        strip(standalone.read_text(encoding="utf-8"))


# ---------- preview opens the right one -------------------------------

def test_preview_prefers_the_viewable_copy(tmp_path: Path):
    from unittest.mock import patch

    from click.testing import CliRunner

    dist = tmp_path / "dist"
    dist.mkdir(parents=True)
    (dist / "issue-5.html").write_text("<html>mail</html>", encoding="utf-8")
    viewable = dist / "issue-5.preview.html"
    viewable.write_text("<html>viewable</html>", encoding="utf-8")

    with patch("webbrowser.open") as opener:
        result = CliRunner().invoke(
            bn.cli, ["preview", "--issue", "5", "--output-dir", str(tmp_path)])

    assert result.exit_code == 0, result.output
    assert opener.call_args.args[0] == viewable.as_uri(), (
        "preview opened the mail artefact, whose photos are unpublished URLs")


def test_preview_still_works_for_builds_made_before_this_change(tmp_path: Path):
    """No `.preview.html` on disk -- an issue built by an older version.
    Falling back keeps `preview` working rather than erroring."""
    from unittest.mock import patch

    from click.testing import CliRunner

    dist = tmp_path / "dist"
    dist.mkdir(parents=True)
    mail = dist / "issue-5.html"
    mail.write_text("<html>mail</html>", encoding="utf-8")

    with patch("webbrowser.open") as opener:
        result = CliRunner().invoke(
            bn.cli, ["preview", "--issue", "5", "--output-dir", str(tmp_path)])

    assert result.exit_code == 0, result.output
    assert opener.call_args.args[0] == mail.as_uri()


def test_a_failed_build_removes_both_stale_files(tmp_path: Path):
    """The CLI deletes stale HTML so an editor cannot double-click last
    quarter's file and send it. That has to cover the viewable copy too,
    which is the one they are now told to open."""
    import docx as docx_lib

    empty = tmp_path / "issue-92.docx"
    docx_lib.Document().save(str(empty))

    dist = tmp_path / "dist"
    dist.mkdir(parents=True, exist_ok=True)
    mail = dist / "issue-92.html"
    standalone = dist / "issue-92.preview.html"
    mail.write_text("<html>last quarter</html>", encoding="utf-8")
    standalone.write_text("<html>last quarter</html>", encoding="utf-8")

    result = bn._build_pipeline(empty, issue=92, validate_remote=False,
                                output_dir=tmp_path)
    assert result.exit_code == 1
    assert not mail.exists(), "stale mail artefact survived"
    assert not standalone.exists(), (
        "stale viewable copy survived -- the editor is told to open this "
        "one, so it is the more dangerous of the two to leave")
