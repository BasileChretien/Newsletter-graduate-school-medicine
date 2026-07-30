"""Tests for `scripts.webapp` -- the browser/server build entry point.

Runs the real production template through the real pipeline, so these
double as an integration check on `parse -> render -> inline ->
validate -> build_eml` with no CLI, no filesystem side effects in the
repo, and no mail client.
"""

from __future__ import annotations

import email
import re
from pathlib import Path

import docx
import pytest

from scripts.config import ASSETS_DIR, DIST_DIR, MERIDIAN_TEMPLATE
from scripts.webapp import (
    WebBuildResult, build_from_bytes, subject_from_masthead,
)

# The masthead placeholders the validator hard-blocks on. Filling them
# is what turns the shipped template into a "valid issue" for testing.
MASTHEAD_FILL = {
    "VOL. XX": "VOL. 2",
    "ISSUE NO. XX": "ISSUE NO. 3",
    "MONTH YEAR": "MARCH 2026",
}


def _fill_paragraph(paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)


def _docx_bytes(tmp_path: Path, replacements: dict[str, str] | None) -> bytes:
    """Copy the production template, optionally filling the masthead."""
    out = tmp_path / "issue.docx"
    d = docx.Document(str(MERIDIAN_TEMPLATE))
    if replacements:
        for p in d.paragraphs:
            _fill_paragraph(p, replacements)
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _fill_paragraph(p, replacements)
    d.save(str(out))
    return out.read_bytes()


@pytest.fixture(scope="module")
def filled_docx(tmp_path_factory) -> bytes:
    return _docx_bytes(tmp_path_factory.mktemp("filled"), MASTHEAD_FILL)


@pytest.fixture(scope="module")
def unfilled_docx(tmp_path_factory) -> bytes:
    return _docx_bytes(tmp_path_factory.mktemp("unfilled"), None)


@pytest.fixture(scope="module")
def built(filled_docx, tmp_path_factory) -> WebBuildResult:
    return build_from_bytes(
        filled_docx, issue=3,
        workdir=tmp_path_factory.mktemp("build"))


def _content_ids(msg) -> set[str]:
    return {p["Content-ID"].strip().lstrip("<").rstrip(">")
            for p in msg.walk() if p["Content-ID"]}


# ---------- the happy path ----------------------------------------------

def test_a_filled_template_builds_successfully(built):
    assert built.ok is True
    assert built.errors == ()
    assert built.section_count == 7
    assert built.eml is not None
    assert built.size_bytes > 10_000


def test_subject_comes_from_the_filled_masthead(built):
    # `sanitize_subject` collapses the template's double spaces.
    assert built.subject == "MERIDIAN — VOL. 2 | ISSUE NO. 3 | MARCH 2026"


def test_the_eml_is_a_valid_draft_with_the_photos_inside(built):
    msg = email.message_from_bytes(built.eml)

    assert msg["X-Unsent"] == "1"
    assert msg.get_content_type() == "multipart/alternative"
    # 2 photos embedded in the DOCX + the masthead logo mirrored in
    # from `images/`. If this drops to 2, brand assets stopped being
    # mirrored and CID mode is shipping a partly-external email.
    assert built.photo_count == 3
    images = [p for p in msg.walk() if p.get_content_type().startswith("image/")]
    assert len(images) == built.photo_count


def test_brand_assets_are_embedded_not_left_as_urls(built):
    """The logo lives in `images/`, not in the DOCX. A build that only
    embeds DOCX media ships a mixed email -- some photos inline, some
    fetched from GitHub -- which is exactly what breaks on the
    hospital networks that quarantine raw.githubusercontent.com."""
    msg = email.message_from_bytes(built.eml)
    filenames = {p.get_filename() for p in msg.walk() if p.get_filename()}

    assert any("logo" in (f or "").lower() for f in filenames), filenames


def test_every_cid_reference_in_the_eml_resolves(built):
    """The failure this guards against is silent: a structurally valid
    .eml whose CIDs don't match renders as broken-image icons for all
    50 recipients, and nobody notices until after the send."""
    msg = email.message_from_bytes(built.eml)
    html_part = next(p for p in msg.walk()
                     if p.get_content_type() == "text/html")
    body = html_part.get_payload(decode=True).decode("utf-8", "replace")

    referenced = set(re.findall(r'src="cid:([^"]+)"', body))
    assert referenced, "CID mode produced no cid: references"
    assert referenced <= _content_ids(msg)


# ---------- preview vs. wire HTML ---------------------------------------

def test_preview_embeds_photos_as_data_uris(built):
    """A brand-new issue has nothing on raw.githubusercontent.com yet,
    so a preview using those URLs would show broken images and read as
    'the toolkit is broken'."""
    assert "data:image/" in built.preview_html
    # >= rather than ==: the logo is referenced twice (masthead and
    # footer), and both references get the same data URI.
    assert built.preview_html.count("data:image/") >= built.photo_count
    assert "raw.githubusercontent.com" not in built.preview_html


def test_wire_html_keeps_the_public_urls(built):
    """`html` must stay byte-comparable with what the CLI writes to
    dist/issue-N.html -- it is the hosted-images artefact."""
    assert "raw.githubusercontent.com" in built.html
    assert "data:image/" not in built.html


# ---------- editorial mistakes are results, not exceptions --------------

def test_unfilled_masthead_is_rejected_without_an_eml(unfilled_docx, tmp_path):
    """The CLI deletes dist/issue-N.html on a hard block so nothing
    sendable survives. The browser equivalent: a preview to look at,
    but no .eml to download."""
    res = build_from_bytes(unfilled_docx, issue=1, workdir=tmp_path)

    assert res.ok is False
    assert res.eml is None
    assert any("VOL. XX" in e or "placeholder" in e.lower()
               for e in res.errors), res.errors
    # The editor still needs to SEE what is wrong.
    assert res.preview_html


def test_an_empty_document_is_rejected_with_readable_advice(tmp_path):
    empty = tmp_path / "empty.docx"
    docx.Document().save(str(empty))

    res = build_from_bytes(empty.read_bytes(), issue=1, workdir=tmp_path)

    assert res.ok is False
    assert res.eml is None
    assert "password-protected" in res.errors[0]
    # No developer jargon in an editor-facing message.
    for token in ("Traceback", "None", "docx_parser"):
        assert token not in res.errors[0]


def test_an_oversized_photo_is_reported_not_silently_linked(
        filled_docx, tmp_path):
    """`attach_inline_images` leaves photos over its 2 MB cap as remote
    URLs. On the CLI that degrades gracefully -- `publish-images` pushes
    the file and the URL resolves. **The browser has no publish step**,
    so the URL points at nothing and every recipient sees a broken
    image, while the result screen reports how many photos were
    embedded. The editor has to be told."""
    res = build_from_bytes(filled_docx, issue=3, workdir=tmp_path)
    assert res.photo_count == 3, "precondition: all photos embed normally"

    # Re-run with a cap low enough that nothing can be embedded.
    capped = build_from_bytes(filled_docx, issue=4,
                              workdir=tmp_path / "capped",
                              max_image_bytes=10)

    assert capped.photo_count == 0
    assert capped.unembedded_photo_count > 0
    assert any("broken image" in w for w in capped.warnings), capped.warnings
    # Still sendable -- one oversized photo must not block the issue.
    assert capped.ok is True


def test_url_mode_does_not_warn_about_unembedded_photos(filled_docx, tmp_path):
    """In URL mode remote images are the whole point, so the warning
    would be noise -- and noise trains editors to ignore warnings."""
    res = build_from_bytes(filled_docx, issue=3, image_mode="url",
                           workdir=tmp_path)

    assert res.unembedded_photo_count == 0
    assert not any("broken image" in w for w in res.warnings)


def test_issue_number_cannot_escape_the_workdir(filled_docx, tmp_path):
    """`issue` reaches the filesystem via `issue_dir()`. A server
    handler wiring a query parameter straight through must not be able
    to write outside the workdir."""
    for bad in ["../../etc", "1; rm -rf /", None, 3.7e300, True, [3]]:
        with pytest.raises(ValueError, match="issue"):
            build_from_bytes(filled_docx, issue=bad, workdir=tmp_path)
    with pytest.raises(ValueError, match="negative"):
        build_from_bytes(filled_docx, issue=-1, workdir=tmp_path)
    with pytest.raises(ValueError, match="at most"):
        # An unbounded value became a 300-digit filename and a bare OSError.
        build_from_bytes(filled_docx, issue=10 ** 40, workdir=tmp_path)


def test_unknown_image_mode_is_a_caller_error(filled_docx, tmp_path):
    """A bad `image_mode` is a programming mistake, not an editorial
    one -- it raises rather than returning a result the UI would have
    to render."""
    with pytest.raises(ValueError, match="image_mode"):
        build_from_bytes(filled_docx, issue=1, image_mode="base64",
                         workdir=tmp_path)


# ---------- url mode ----------------------------------------------------

def test_url_mode_ships_no_attachments(filled_docx, tmp_path):
    res = build_from_bytes(filled_docx, issue=3, image_mode="url",
                           workdir=tmp_path)
    msg = email.message_from_bytes(res.eml)

    assert res.ok is True
    assert res.photo_count == 0
    assert not [p for p in msg.walk()
                if p.get_content_type().startswith("image/")]
    body = next(p for p in msg.walk()
                if p.get_content_type() == "text/html")
    decoded = body.get_payload(decode=True).decode("utf-8", "replace")
    assert "raw.githubusercontent.com" in decoded
    assert "cid:" not in decoded


# ---------- recipients --------------------------------------------------

def test_bcc_reaches_the_draft_and_is_comma_separated(filled_docx, tmp_path):
    """`recipients.txt` is joined with "; " for Outlook COM; RFC 5322
    needs commas or every address after the first is dropped."""
    res = build_from_bytes(
        filled_docx, issue=3, workdir=tmp_path,
        bcc="a@example.ac.jp; b@example.ac.jp; c@example.ac.jp")
    msg = email.message_from_bytes(res.eml)

    assert msg["Bcc"] == "a@example.ac.jp, b@example.ac.jp, c@example.ac.jp"
    assert res.recipient_count == 3


def test_recipients_are_accepted_in_whatever_shape_they_are_pasted(
        filled_docx, tmp_path):
    """Editors paste from Word, Excel and Outlook. Newlines, commas and
    semicolons must all work, and duplicates must collapse."""
    res = build_from_bytes(
        filled_docx, issue=3, workdir=tmp_path,
        bcc="a@example.ac.jp\nb@example.ac.jp, c@example.ac.jp; a@example.ac.jp")
    msg = email.message_from_bytes(res.eml)

    assert msg["Bcc"] == "a@example.ac.jp, b@example.ac.jp, c@example.ac.jp"
    assert res.rejected_addresses == ()


def test_a_newline_in_an_address_cannot_crash_the_build(filled_docx, tmp_path):
    """`EmailMessage` raises ValueError on CR/LF in a header value. Before
    the shared guard, that reached the editor as an unexplained crash
    instead of a message about the address they typed."""
    res = build_from_bytes(
        filled_docx, issue=3, workdir=tmp_path,
        bcc="good@example.ac.jp\r\nBcc: someone-else@evil.test")

    assert res.ok is True
    msg = email.message_from_bytes(res.eml)
    assert "evil.test" not in (msg["Bcc"] or "")
    assert msg["Bcc"] == "good@example.ac.jp"


def test_invalid_addresses_are_dropped_and_reported(filled_docx, tmp_path):
    """Silently dropping them would be worse than including them: the
    editor must find out before they press Send, not after."""
    res = build_from_bytes(
        filled_docx, issue=3, workdir=tmp_path,
        bcc="good@example.ac.jp\nnot an address\nalso-bad@\n")
    msg = email.message_from_bytes(res.eml)

    assert msg["Bcc"] == "good@example.ac.jp"
    assert set(res.rejected_addresses) == {"not an address", "also-bad@"}
    assert any("did not look like" in w for w in res.warnings), res.warnings
    assert res.recipient_count == 1


def test_display_name_recipients_are_rejected_not_silently_truncated(
        filled_docx, tmp_path):
    """`Name <a@x>; Name2 <b@y>` would survive as a single opaque string
    and RFC 5322 would then read the `;` as a group terminator, keeping
    only the first entry. Rejecting is honest; truncating is not."""
    res = build_from_bytes(
        filled_docx, issue=3, workdir=tmp_path,
        bcc="Katsuno <dean@example.ac.jp>; Office <office@example.ac.jp>")

    assert res.recipient_count == 0
    assert len(res.rejected_addresses) == 2
    assert any("did not look like" in w for w in res.warnings)


def test_no_bcc_leaves_the_header_absent(filled_docx, tmp_path):
    res = build_from_bytes(filled_docx, issue=3, workdir=tmp_path, bcc="")
    msg = email.message_from_bytes(res.eml)

    assert msg["Bcc"] is None
    assert res.recipient_count == 0


# ---------- isolation ---------------------------------------------------

def test_nothing_is_written_into_the_repo(filled_docx, tmp_path):
    """The browser build must be self-contained in its workdir. If it
    reached for `ASSETS_DIR` / `DIST_DIR` it would work on a dev
    machine and fail in Pyodide -- or worse, quietly write into a
    fork's checkout."""
    before_assets = set(ASSETS_DIR.rglob("*")) if ASSETS_DIR.exists() else set()
    before_dist = set(DIST_DIR.rglob("*")) if DIST_DIR.exists() else set()

    build_from_bytes(filled_docx, issue=999, workdir=tmp_path)

    after_assets = set(ASSETS_DIR.rglob("*")) if ASSETS_DIR.exists() else set()
    after_dist = set(DIST_DIR.rglob("*")) if DIST_DIR.exists() else set()
    assert after_assets == before_assets
    assert after_dist == before_dist
    # And the photos DID land somewhere -- in the workdir.
    assert (tmp_path / "assets" / "issue-999").is_dir()


def test_workdir_defaults_to_a_temp_directory(filled_docx):
    """Callable with no workdir at all -- the browser passes none."""
    res = build_from_bytes(filled_docx, issue=4)
    assert res.ok is True


# ---------- shared-subject invariant ------------------------------------

def test_cli_and_web_agree_on_the_subject():
    """`build_newsletter._subject_from_masthead` delegates here. If
    someone re-implements it, the manifest's recorded subject and the
    one recipients see would drift apart."""
    import build_newsletter as bn
    from scripts.docx_parser import Masthead

    masthead = Masthead(
        title="MERIDIAN", tagline="Where medicine meets the world.",
        subtitle="Newsletter", issue_line="VOL. 9 | ISSUE NO. 1 | JULY 2026")
    assert bn._subject_from_masthead(9, masthead) == \
        subject_from_masthead(9, masthead)


def test_subject_falls_back_when_the_masthead_is_missing():
    assert subject_from_masthead(12, None) == "MERIDIAN — Issue 12"


def test_the_default_workdir_does_not_outlive_the_call(filled_docx):
    """`build_from_bytes` promises "nothing on disk outlives the call"
    when it creates its own workdir. The previous test asserted only
    `ok is True` -- deleting the entire try/finally left it passing.

    This branch has no production coverage either: the browser passes an
    explicit `workdir`, so nothing else exercises the cleanup at all."""
    import glob
    import os
    import tempfile

    pattern = os.path.join(tempfile.gettempdir(), "meridian-web-*")
    before = set(glob.glob(pattern))

    res = build_from_bytes(filled_docx, issue=4)

    assert res.ok is True
    assert set(glob.glob(pattern)) == before, "temp workdir was left behind"


# ---------- CLI / web parity --------------------------------------------

def test_the_cli_and_the_web_build_produce_identical_html(
        filled_docx, tmp_path, monkeypatch):
    """The two pipelines duplicate `parse -> attach URLs -> render ->
    inline`, and they have already drifted once inside a single branch
    (brand assets, the over-cap warning). Byte-equality is assertable
    because both derive URLs from `assets/issue-N/<file>` relative to
    their own root, so a divergence in ANY of those four steps fails
    here rather than in an editor's inbox."""
    import build_newsletter as bn

    docx_path = tmp_path / "issue-3.docx"
    docx_path.write_bytes(filled_docx)

    # Redirect the CLI's output roots by monkeypatching rather than via
    # `--output-dir`: that flag currently raises from `to_raw_url`
    # (`Asset ... must be inside repo ...`) for any DOCX containing
    # photos, which is being fixed separately. Patching PROJECT_ROOT too
    # keeps the URL root consistent with the redirected assets, so this
    # test exercises the parity property without depending on that fix
    # -- and without writing into the real repo.
    cli_out = tmp_path / "cli"
    monkeypatch.setattr(bn, "DIST_DIR", cli_out / "dist")
    monkeypatch.setattr(bn, "ASSETS_DIR", cli_out / "assets")
    monkeypatch.setattr(bn, "PROJECT_ROOT", cli_out)

    result = bn._build_pipeline(docx_path, 3, validate_remote=False)
    assert result.exit_code == 0, "precondition: the CLI build succeeds"
    cli_html = result.html_path.read_text(encoding="utf-8")

    web = build_from_bytes(filled_docx, issue=3, workdir=tmp_path / "web")
    assert web.ok is True

    assert web.html == cli_html, (
        "the CLI and browser pipelines produced different HTML -- one of "
        "the shared render steps has drifted"
    )


def test_webapp_imports_without_os_specific_modules(monkeypatch):
    """`scripts.webapp` transitively imports the whole `scripts.mail`
    package, which reaches Outlook COM, the clipboard and handler
    detection. That only works in Pyodide because every OS-specific
    import is function-local -- nothing enforces it, and CI runs no
    Pyodide job, so a future backend with a module-level
    `import win32com.client` would break the browser build while every
    existing test stayed green (the bundle would be current, just
    broken)."""
    import importlib
    import sys

    blocked = ("win32com", "win32com.client", "winreg", "win32clipboard",
               "pywintypes", "tkinter")
    for name in blocked:
        monkeypatch.setitem(sys.modules, name, None)   # import -> ImportError

    for mod in ("scripts.webapp", "scripts.mail", "scripts.mail.eml",
                "scripts.subject", "scripts.recipients"):
        sys.modules.pop(mod, None)
    try:
        importlib.import_module("scripts.webapp")
    except ImportError as e:
        raise AssertionError(
            f"scripts.webapp needs an OS-specific module at import time: {e}. "
            "Move that import inside the function that uses it."
        ) from e
