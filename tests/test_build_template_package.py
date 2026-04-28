"""Smoke tests for the `scripts.build_template` package split (bundle 27).

The 600-line module was split into a package with `_styles.py` +
`_elements.py` + `__init__.py`. These tests pin the public surface so
a future refactor can't silently break the call sites that depend on
`from scripts import build_template as bt`.
"""

from __future__ import annotations

import pytest


def test_build_function_is_importable_via_package():
    """`from scripts import build_template as bt; bt.build` -- the
    everyday smoke path. Must keep working after the split."""
    from scripts import build_template as bt
    assert callable(bt.build)


def test_table_constants_are_exposed():
    """Named table indices were on the top-level module before the
    split; they must still be reachable via `bt.TABLE_*` so any
    future maintenance reading the orchestration in `build()` finds
    them where they expect."""
    from scripts import build_template as bt
    assert bt.TABLE_MASTHEAD == 0
    assert bt.TABLE_DEAN == 1
    assert bt.TABLE_HIGHLIGHTS_TOP == 2
    assert bt.TABLE_HIGHLIGHTS_BOTTOM == 3
    assert bt.TABLE_VISITORS == 4
    assert bt.TABLE_EVENTS == 5
    assert bt.TABLE_CONTACT == 6


def test_palette_constants_re_exported():
    """`bt.PRIMARY` / `ACCENT` / `TEXT` / `MUTED` were used by ad-hoc
    debug shells; keep them reachable on the package."""
    from scripts import build_template as bt
    assert bt.PRIMARY is not None
    assert bt.ACCENT is not None
    assert bt.TEXT is not None
    assert bt.MUTED is not None


def test_helper_functions_re_exported():
    """`is_section_heading` is the canonical "is this a numbered
    section heading?" predicate; pin it on the package."""
    from scripts import build_template as bt
    assert callable(bt.is_section_heading)
    assert bt.is_section_heading("1. Hospital News")
    assert not bt.is_section_heading("This is body text.")


def test_styling_helpers_re_exported():
    """style_run / style_paragraph / rgb are public helpers.
    `_normalize_body_run` keeps its leading underscore (private-to-package)
    and is intentionally NOT in `__all__` -- bundle 28 cleanup
    (round-9 python-reviewer MEDIUM)."""
    from scripts import build_template as bt
    assert callable(bt.style_run)
    assert callable(bt.style_paragraph)
    assert callable(bt.rgb)
    # _normalize_body_run is reachable but NOT public.
    assert callable(bt._normalize_body_run)
    assert "_normalize_body_run" not in bt.__all__, (
        "Leading-underscore names must not appear in __all__."
    )


def test_restyle_functions_re_exported():
    """All restyle_* / insert_* are on the package for parity with
    the pre-split call surface."""
    from scripts import build_template as bt
    expected = [
        "restyle_masthead", "restyle_section_heading", "restyle_subhead",
        "restyle_body", "restyle_bullet",
        "restyle_data_table", "restyle_layout_table",
        "restyle_highlights_table", "restyle_header_footer",
        "insert_dean_name", "insert_dean_photo",
        "configure_page",
    ]
    for name in expected:
        assert hasattr(bt, name), f"build_template lost public name: {name}"
        assert callable(getattr(bt, name))


def test_section_head_re_pattern_unchanged():
    """The section-heading regex hasn't changed during the split."""
    from scripts import build_template as bt
    assert bt.SECTION_HEAD_RE.match("1.  Some Section") is not None
    assert bt.SECTION_HEAD_RE.match("12. Twelfth Title") is not None
    assert bt.SECTION_HEAD_RE.match("Just body text.") is None


def test_submodules_importable_directly():
    """Bundle 28 dropped the leading underscore on the submodule
    names (`styles.py` / `elements.py`); the package boundary
    already encapsulates them. Import paths used by debug shells:"""
    from scripts.build_template import styles, elements
    assert hasattr(styles, "rgb")
    assert hasattr(elements, "is_section_heading")


def test_underscore_submodules_are_gone():
    """Round-9 architect MEDIUM 2: the legacy `_styles` / `_elements`
    module names should no longer exist after the rename. If they
    reappear, someone has reverted the bundle-28 cleanup."""
    import importlib
    for name in ("scripts.build_template._styles",
                 "scripts.build_template._elements"):
        try:
            importlib.import_module(name)
        except ImportError:
            continue
        else:
            raise AssertionError(
                f"{name} should not exist post-bundle-28 (use the "
                "underscore-free name)."
            )


def test_style_run_actually_styles_a_run():
    """Behaviour smoke check (round-9 code-review MEDIUM): the
    `callable(bt.style_run)` test would pass even if the function
    were stubbed to do nothing. This test verifies it actually
    applies font / size / colour to a real python-docx run."""
    from scripts import build_template as bt
    import docx
    doc = docx.Document()
    p = doc.add_paragraph()
    run = p.add_run("hello")
    bt.style_run(run, font="Cambria", size_pt=14, bold=True,
                 color=bt.PRIMARY)
    assert run.font.name == "Cambria"
    assert run.font.size.pt == 14
    assert run.bold is True
    # Colour comparison via the RGBColor's underlying tuple.
    assert run.font.color.rgb == bt.PRIMARY


def test_restyle_section_heading_smoke():
    """Behaviour smoke check (round-9 code-review MEDIUM): hand the
    restyler a real `1.  Title` paragraph and confirm it produces
    runs with the expected colours / sizes."""
    from scripts import build_template as bt
    import docx
    doc = docx.Document()
    p = doc.add_paragraph("1.  Hospital News")
    bt.restyle_section_heading(p)
    # After restyling, the paragraph should have at least 3 runs:
    # numeral, gold dash, label.
    assert len(p.runs) >= 3
    # The first run carries the zero-padded numeral in NU blue.
    assert "01" in p.runs[0].text
    assert p.runs[0].font.color.rgb == bt.PRIMARY
    # One of the runs is the gold dash.
    dash_runs = [r for r in p.runs if r.font.color.rgb == bt.ACCENT]
    assert dash_runs, "Section heading must carry a gold-coloured dash run"


def test_build_template_module_does_not_exist_as_file():
    """Bundle 27 deleted the old `scripts/build_template.py` and
    replaced it with a package. Confirm the file is gone so a stale
    `.py` can't shadow the package on PYTHONPATH."""
    import scripts.build_template as bt
    # __file__ should point at the package's __init__.py, not a .py.
    assert bt.__file__.endswith("__init__.py")
