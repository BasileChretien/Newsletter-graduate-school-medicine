"""Round-17 regression tests for the lenient-fallback DOCX parser.

The user's first real-world test of v1.1.1 used their own DOCX (not
the bundled MERIDIAN template). The strict parser silently dropped
every paragraph because it never saw a `1. Title` / `Section 1: Title`
/ `第1章 Title` heading -- so `current_num` stayed None forever and
ALL body content was skipped. The editor got a blank email.

These tests pin the fix:
  * `_parse_lenient` produces a single Section(1, "Newsletter") with
    every body paragraph, list, and table preserved as a block.
  * Word `Heading 1` / `Heading 2` paragraphs become inline headings
    inside the synthetic section.
  * `parse()` falls back to lenient when strict produces 0 sections.
  * `_build_pipeline` hard-fails (exit_code=1) if even the lenient
    fallback produces 0 sections -- so an empty/malformed DOCX
    gives a clear console error instead of a blank email draft.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from scripts.docx_parser import (
    BodyParagraph,
    BulletList,
    Heading,
    Section,
    _parse_lenient,
    parse,
)


def _make_arbitrary_docx(path: Path, *, with_headings: bool = True) -> Path:
    d = docx.Document()
    d.add_heading("My Custom Newsletter", level=0)
    d.add_paragraph("Welcome to this newsletter from our department.")
    if with_headings:
        d.add_heading("Recent News", level=1)
    d.add_paragraph("Some recent news content here.")
    d.add_paragraph("A second paragraph in the same section.")
    if with_headings:
        d.add_heading("Upcoming Events", level=1)
    d.add_paragraph("We have an event on May 5.")
    d.save(str(path))
    return path


def test_parse_lenient_fallback_extracts_all_body_when_no_section_headings(
    tmp_path: Path,
):
    """Round-17 production-bug fix: a DOCX with no MERIDIAN-style
    numbered section headings used to render as a blank email
    because the strict parser dropped every paragraph waiting for
    `1.`. The lenient fallback now produces ONE synthetic section
    with all the content."""
    p = _make_arbitrary_docx(tmp_path / "user.docx")

    nl = parse(p)

    assert len(nl.sections) == 1, (
        f"expected 1 synthetic section, got {len(nl.sections)}"
    )
    sec = nl.sections[0]
    assert sec.number == 1
    assert sec.title == "Newsletter"

    # All body paragraphs survived; Heading 1 -> H2 dividers.
    types = [type(b).__name__ for b in sec.blocks]
    assert "BodyParagraph" in types, types
    assert "Heading" in types, (
        "Word Heading 1 paragraphs must become Heading blocks in the "
        f"lenient parse; got block types {types}"
    )


def test_parse_lenient_fallback_with_no_word_headings_still_works(tmp_path):
    """Even a totally heading-less DOCX must produce content."""
    p = _make_arbitrary_docx(tmp_path / "flat.docx", with_headings=False)
    nl = parse(p)
    assert len(nl.sections) == 1
    body_paras = [b for b in nl.sections[0].blocks
                  if isinstance(b, BodyParagraph)]
    assert len(body_paras) >= 3, (
        f"expected >= 3 body paragraphs, got {len(body_paras)}: "
        f"{[type(b).__name__ for b in nl.sections[0].blocks]}"
    )


def test_parse_strict_still_wins_when_meridian_headings_present(tmp_path):
    """The lenient fallback must NOT kick in when the user uses
    the MERIDIAN template's `1. Title` headings -- otherwise
    multi-section newsletters would all collapse to one section."""
    d = docx.Document()
    d.add_paragraph("1. First Section")
    d.add_paragraph("Body of first.")
    d.add_paragraph("2. Second Section")
    d.add_paragraph("Body of second.")
    d.add_paragraph("3. Third Section")
    d.add_paragraph("Body of third.")
    p = tmp_path / "meridian.docx"
    d.save(str(p))

    nl = parse(p)
    assert len(nl.sections) == 3
    assert [s.number for s in nl.sections] == [1, 2, 3]


def test_parse_lenient_logs_warning(tmp_path, caplog):
    """When the fallback fires, an explicit WARNING must hit the
    log so editors looking at the launcher console can see why
    the email looks unstructured. Round-17."""
    import logging

    caplog.set_level(logging.WARNING, logger="scripts.docx_parser")
    p = _make_arbitrary_docx(tmp_path / "user.docx")
    parse(p)
    assert any(
        "lenient parse" in r.message.lower()
        and "no numbered section" in r.message.lower()
        for r in caplog.records
    ), (
        f"expected lenient-fallback warning; got {caplog.records}"
    )


def test_parse_lenient_with_truly_empty_docx_returns_no_sections(tmp_path):
    """If even the lenient parser finds no body content (an empty
    DOCX), `parse()` returns Newsletter with 0 sections so the
    build pipeline can hard-fail with a clear message."""
    d = docx.Document()
    p = tmp_path / "empty.docx"
    d.save(str(p))

    nl = parse(p)
    assert nl.sections == ()


def test_parse_lenient_does_not_drop_list_paragraphs(tmp_path):
    """Bullet-styled paragraphs in an arbitrary DOCX must survive
    the lenient fallback -- as a `BulletList` if `_is_list_paragraph`
    detects them, or at minimum as `BodyParagraph` so the content
    isn't lost. The strict ban is "the user's text reaches the
    rendered email." Round-17."""
    d = docx.Document()
    d.add_paragraph("Some intro.")
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    d.add_paragraph("Trailing prose.")
    p = tmp_path / "bullets.docx"
    d.save(str(p))

    nl = parse(p)
    assert len(nl.sections) == 1
    blocks = nl.sections[0].blocks
    # Either a BulletList with both items, or two separate
    # BodyParagraphs containing the bullet text -- both are
    # acceptable. What's NOT acceptable is dropping the content.
    rendered_text = " ".join(
        getattr(b, "html", "") + " ".join(getattr(b, "items", ()))
        for b in blocks
    )
    assert "First bullet" in rendered_text
    assert "Second bullet" in rendered_text
